from __future__ import annotations

from datetime import date
from io import BytesIO
from pathlib import Path
import base64
import html
import re
import textwrap

import pandas as pd
import plotly.express as px
import streamlit as st
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# --------------------------------------------------
# PAGE SETUP
# --------------------------------------------------
HSD_WEBSITE = "https://hsdmetrics.com"
HSD_LOGO_URL = "https://mma.prnewswire.com/media/2925136/HSD_Logo_Dark_Blue_Logo.jpg"
HSD_LOGO_PATH = Path(__file__).with_name("hsd_logo.jpg")


def get_logo_src() -> str:
    """Use the local HSD logo when available, with a remote fallback."""
    if HSD_LOGO_PATH.exists():
        encoded = base64.b64encode(HSD_LOGO_PATH.read_bytes()).decode("ascii")
        return f"data:image/jpeg;base64,{encoded}"
    return HSD_LOGO_URL


HSD_LOGO_SRC = get_logo_src()

st.set_page_config(
    page_title="HSD HQ Brief",
    page_icon="📊",
    layout="wide",
)

# --------------------------------------------------
# HSD STYLE COLORS
# --------------------------------------------------
HSD_NAVY = "#1B2A4A"
HSD_BLUE = "#2D6DB5"
HSD_MEDIUM_BLUE = "#4A90D9"
HSD_LIGHT_BLUE = "#EBF4FF"
HSD_SKY_BLUE = "#BFDFFF"
HSD_BG = "#F8FAFC"
HSD_BG2 = "#F1F5F9"
HSD_TEXT = "#1A202C"
HSD_MUTED = "#6B7280"
HSD_BORDER = "#E2E8F0"
HSD_WHITE = "#FFFFFF"

BLUE_SCALE = [HSD_NAVY, HSD_BLUE, HSD_MEDIUM_BLUE, HSD_SKY_BLUE]

# Public-facing directional estimates by employee-count band.
# These ranges intentionally avoid displaying exact internal pricing.
PYTHIA_ESTIMATED_PRICING = {
    "Under 500": {"annual": (4500, 5500), "setup": (2000, 3000)},
    "500 to 749": {"annual": (5500, 6500), "setup": (3000, 4000)},
    "750 to 999": {"annual": (5500, 6500), "setup": (3000, 4000)},
    "1,000 to 1,499": {"annual": (5500, 6500), "setup": (4000, 5000)},
    "1,500 to 1,999": {"annual": (5500, 6500), "setup": (4000, 5000)},
    "2,000 to 2,499": {"annual": (5500, 6500), "setup": (4000, 5000)},
    "2,500 to 2,999": {"annual": (6500, 7500), "setup": (4000, 5000)},
    "3,000 to 3,499": {"annual": (6500, 7500), "setup": (4000, 5000)},
    "3,500 to 3,999": {"annual": (6500, 7500), "setup": (4000, 5000)},
    "4,000 to 4,499": {"annual": (6500, 7500), "setup": (4000, 5000)},
    "4,500 to 4,999": {"annual": (6500, 7500), "setup": (4500, 5500)},
    "5,000 to 5,999": {"annual": (8500, 9500), "setup": (4500, 5500)},
    "6,000 to 6,999": {"annual": (8500, 9500), "setup": (4500, 5500)},
    "7,000 to 7,999": {"annual": (8500, 9500), "setup": (4500, 5500)},
    "8,000 to 8,999": {"annual": (8500, 9500), "setup": (4500, 5500)},
    "9,000 to 9,999": {"annual": (8500, 9500), "setup": (4500, 5500)},
    "10,000 to 12,499": {"annual": (9500, 10500), "setup": (5500, 6500)},
    "12,500 to 14,999": {"annual": (9500, 10500), "setup": (5500, 6500)},
    "15,000 to 19,999": {"annual": (9500, 10500), "setup": (5500, 6500)},
    "20,000 to 24,999": {"annual": (9500, 10500), "setup": (5500, 6500)},
    "25,000 to 29,999": {"annual": (9500, 10500), "setup": (5500, 6500)},
    "30,000 to 34,999": {"annual": (9500, 10500), "setup": (5500, 6500)},
    "35,000 to 39,999": {"annual": (9500, 10500), "setup": (5500, 6500)},
}

# --------------------------------------------------
# CUSTOM CSS
# --------------------------------------------------
st.markdown(
    f"""
    <style>
    .stApp {{
        background-color: {HSD_BG};
        color: {HSD_TEXT};
    }}

    section[data-testid="stSidebar"] {{
        background-color: {HSD_BG2};
        border-right: 1px solid {HSD_BORDER};
    }}

    section[data-testid="stSidebar"] > div {{
        padding-bottom: 110px;
    }}

    h1, h2, h3 {{
        color: {HSD_NAVY};
        font-weight: 700;
    }}

    div[data-testid="stMetric"] {{
        background: {HSD_WHITE};
        border: 1px solid {HSD_BORDER};
        border-left: 6px solid {HSD_BLUE};
        border-radius: 14px;
        padding: 18px;
        box-shadow: 0 2px 8px rgba(27, 42, 74, 0.08);
    }}

    div[data-testid="stMetricLabel"] {{
        color: {HSD_MUTED};
        font-size: 13px;
    }}

    div[data-testid="stMetricValue"] {{
        color: {HSD_NAVY};
        font-weight: 700;
    }}

    .hsd-header {{
        display: flex;
        align-items: center;
        gap: 26px;
        background: linear-gradient(135deg, {HSD_NAVY}, {HSD_BLUE});
        padding: 24px 28px;
        border-radius: 18px;
        margin-bottom: 18px;
        color: white;
    }}

    .hsd-logo-wrap {{
        background: white;
        padding: 10px 14px;
        border-radius: 12px;
        min-width: 175px;
        text-align: center;
    }}

    .hsd-logo-wrap img {{
        width: 155px;
        max-width: 100%;
        display: block;
    }}

    .hsd-header h1 {{
        color: white;
        margin: 0 0 6px 0;
        font-size: 34px;
    }}

    .hsd-header p {{
        color: rgba(255, 255, 255, 0.90);
        font-size: 16px;
        margin: 0;
    }}

    .hsd-results {{
        background: {HSD_WHITE};
        border: 1px solid {HSD_BORDER};
        border-radius: 16px;
        padding: 18px 20px;
        margin: 18px 0 22px 0;
        box-shadow: 0 2px 10px rgba(27, 42, 74, 0.08);
    }}

    .hsd-results-title {{
        color: {HSD_NAVY};
        font-size: 17px;
        font-weight: 700;
        margin-bottom: 14px;
    }}

    .hsd-results-grid {{
        display: grid;
        grid-template-columns: 1fr 1fr 1fr 1.65fr;
        align-items: stretch;
        gap: 0;
    }}

    .hsd-result-stat {{
        display: flex;
        align-items: center;
        gap: 12px;
        padding: 8px 18px 8px 8px;
        border-right: 1px solid {HSD_BORDER};
        min-height: 86px;
    }}

    .hsd-result-icon {{
        width: 48px;
        height: 48px;
        min-width: 48px;
        border-radius: 50%;
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 23px;
        font-weight: 700;
    }}

    .hsd-icon-blue {{
        background: #E8F1FF;
        color: #1769E0;
    }}

    .hsd-icon-green {{
        background: #E6F7F2;
        color: #07946F;
    }}

    .hsd-icon-purple {{
        background: #F0EAFE;
        color: #7047D9;
    }}

    .hsd-result-value {{
        font-size: 26px;
        line-height: 1;
        font-weight: 800;
        margin-bottom: 7px;
    }}

    .hsd-value-blue {{
        color: #1769E0;
    }}

    .hsd-value-green {{
        color: #07946F;
    }}

    .hsd-value-purple {{
        color: #7047D9;
    }}

    .hsd-result-label {{
        color: {HSD_TEXT};
        font-size: 13px;
        line-height: 1.35;
    }}

    .hsd-result-quote {{
        background: #EEF6FF;
        border-radius: 13px;
        margin-left: 18px;
        padding: 15px 18px;
        color: {HSD_NAVY};
        min-height: 86px;
    }}

    .hsd-quote-mark {{
        color: {HSD_MEDIUM_BLUE};
        font-size: 32px;
        font-weight: 800;
        line-height: 0.7;
        margin-bottom: 5px;
    }}

    .hsd-quote-text {{
        font-size: 13px;
        line-height: 1.45;
        margin-bottom: 7px;
    }}

    .hsd-quote-source {{
        font-size: 12px;
        font-weight: 700;
        text-align: right;
    }}

    .hsd-card {{
        background: {HSD_WHITE};
        border: 1px solid {HSD_BORDER};
        border-radius: 14px;
        padding: 18px;
        box-shadow: 0 2px 8px rgba(27, 42, 74, 0.08);
        margin-bottom: 16px;
    }}

    .hsd-blue-box {{
        background: {HSD_LIGHT_BLUE};
        border: 1px solid {HSD_SKY_BLUE};
        border-radius: 14px;
        padding: 18px;
        color: {HSD_NAVY};
        margin-top: 12px;
    }}

    .hsd-prospect-line {{
        background: {HSD_WHITE};
        border: 1px solid {HSD_BORDER};
        border-radius: 14px;
        padding: 15px 18px;
        margin-bottom: 16px;
        box-shadow: 0 2px 8px rgba(27, 42, 74, 0.06);
        color: {HSD_NAVY};
        font-size: 15px;
    }}

    .hsd-prospect-line strong {{
        font-size: 19px;
    }}

    .hsd-insight-box {{
        background: linear-gradient(135deg, #F4F8FF, #EBF4FF);
        border: 1px solid {HSD_SKY_BLUE};
        border-left: 6px solid {HSD_BLUE};
        border-radius: 14px;
        padding: 18px 20px;
        margin-top: 16px;
        color: {HSD_TEXT};
    }}

    .hsd-insight-box h3 {{
        margin-top: 0;
        margin-bottom: 10px;
        color: {HSD_NAVY};
    }}

    .hsd-insight-box p {{
        margin: 7px 0;
        line-height: 1.5;
    }}

    .hsd-pricing-box {{
        background: {HSD_WHITE};
        border: 1px solid {HSD_BORDER};
        border-radius: 14px;
        padding: 18px;
        box-shadow: 0 2px 8px rgba(27, 42, 74, 0.06);
        min-height: 180px;
    }}

    .hsd-pricing-row {{
        display: flex;
        justify-content: space-between;
        gap: 18px;
        padding: 10px 0;
        border-bottom: 1px solid {HSD_BORDER};
    }}

    .hsd-pricing-row:last-child {{
        border-bottom: none;
    }}

    .hsd-pricing-label {{
        color: {HSD_MUTED};
    }}

    .hsd-pricing-value {{
        color: {HSD_NAVY};
        font-weight: 700;
        text-align: right;
    }}

    .stTabs [data-baseweb="tab-list"] {{
        gap: 18px;
        border-bottom: 1px solid {HSD_BORDER};
    }}

    .stTabs [data-baseweb="tab"] {{
        color: {HSD_MUTED};
        font-weight: 600;
    }}

    .stTabs [aria-selected="true"] {{
        color: {HSD_BLUE};
        border-bottom: 3px solid {HSD_BLUE};
    }}

    @media (max-width: 1050px) {{
        .hsd-results-grid {{
            grid-template-columns: 1fr 1fr;
            gap: 12px;
        }}

        .hsd-result-stat {{
            border-right: none;
            border-bottom: 1px solid {HSD_BORDER};
        }}

        .hsd-result-quote {{
            grid-column: 1 / -1;
            margin-left: 0;
        }}
    }}

    @media (max-width: 800px) {{
        .hsd-header {{
            display: block;
        }}
        .hsd-logo-wrap {{
            width: fit-content;
            margin-bottom: 16px;
        }}

        .hsd-results-grid {{
            grid-template-columns: 1fr;
        }}

        .hsd-result-quote {{
            grid-column: auto;
        }}
    }}
    </style>
    """,
    unsafe_allow_html=True,
)

# --------------------------------------------------
# HELPER FUNCTIONS
# --------------------------------------------------
def money(value: float) -> str:
    return f"${value:,.0f}"


def compact_money(value: float) -> str:
    """Use compact values for dashboard cards while keeping full values elsewhere."""
    absolute = abs(value)
    sign = "-" if value < 0 else ""

    if absolute >= 1_000_000_000:
        formatted = f"{absolute / 1_000_000_000:.1f}".rstrip("0").rstrip(".")
        return f"{sign}${formatted}B"
    if absolute >= 1_000_000:
        formatted = f"{absolute / 1_000_000:.1f}".rstrip("0").rstrip(".")
        return f"{sign}${formatted}M"
    if absolute >= 1_000:
        formatted = f"{absolute / 1_000:.1f}".rstrip("0").rstrip(".")
        return f"{sign}${formatted}K"
    return f"{sign}${absolute:,.0f}"


def signed_money(value: float) -> str:
    if abs(value) < 0.01:
        return "$0"
    sign = "+" if value > 0 else "-"
    return f"{sign}${abs(value):,.0f}"


def percent(value: float) -> str:
    return f"{value:,.1f}%"


def money_range(low: float, high: float) -> str:
    return money(low) if abs(low - high) < 0.01 else f"{money(low)} - {money(high)}"


def metric_money_range(low: float, high: float) -> str:
    """Use a compact range that fits safely inside Streamlit metric cards."""
    if abs(low - high) < 0.01:
        return compact_money(low)

    low_text = compact_money(low).replace("$", "")
    high_text = compact_money(high).replace("$", "")
    return f"USD {low_text}–{high_text}"


def signed_money_value(value: float) -> str:
    """Format positive savings and negative additional cost clearly."""
    if value < 0:
        return f"-${abs(value):,.0f}"
    return money(value)


def signed_money_range(low: float, high: float) -> str:
    low_value, high_value = sorted([low, high])
    if abs(low_value - high_value) < 0.01:
        return signed_money_value(low_value)
    return f"{signed_money_value(low_value)} - {signed_money_value(high_value)}"


def metric_signed_money_range(low: float, high: float) -> str:
    low_value, high_value = sorted([low, high])
    if abs(low_value - high_value) < 0.01:
        return compact_money(low_value)
    return f"{compact_money(low_value)} to {compact_money(high_value)}"


def apply_hsd_theme(fig):
    fig.update_layout(
        template="plotly_white",
        font=dict(family="Arial", color=HSD_TEXT),
        title_font=dict(color=HSD_NAVY, size=18),
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        margin=dict(l=30, r=30, t=60, b=30),
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=-0.25,
            xanchor="center",
            x=0.5,
        ),
    )
    return fig


def safe_filename(text: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9 _-]+", "", text).strip()
    return cleaned or "Prospect"


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_doc_heading(doc: Document, text: str, size: int = 15) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(27, 42, 74)


def add_doc_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9)


def _doc_font(size: int, bold: bool = False):
    """Load a dependable font for report chart images."""
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def _compact_currency(value: float) -> str:
    absolute = abs(value)
    sign = "-" if value < 0 else ""
    if absolute >= 1_000_000:
        number = f"{absolute / 1_000_000:.1f}".rstrip("0").rstrip(".")
        return f"{sign}${number}M"
    if absolute >= 1_000:
        number = f"{absolute / 1_000:.1f}".rstrip("0").rstrip(".")
        return f"{sign}${number}K"
    return f"{sign}${absolute:,.0f}"


def _compact_range(low: float, high: float) -> str:
    if abs(low - high) < 0.01:
        return _compact_currency(low)
    return f"{_compact_currency(low)}–{_compact_currency(high)}"


def _draw_rounded_bar(draw, xy, fill, radius=12):
    draw.rounded_rectangle(xy, radius=radius, fill=fill)


def _make_current_cost_chart(
    *,
    cost_per_departure: float,
    annual_turnover_cost: float,
    software_cost: float,
    internal_cost: float,
    external_cost: float,
    current_listening_cost: float,
    current_cost_exposure: float,
) -> BytesIO:
    """Create a compact two-scale image for the Word report."""
    width, height = 1500, 820
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    navy = "#1B2A4A"
    blue = "#2D6DB5"
    medium = "#4A90D9"
    sky = "#AFCFF3"
    border = "#D8E2EE"
    muted = "#667085"

    title_font = _doc_font(44, True)
    section_font = _doc_font(29, True)
    label_font = _doc_font(24, False)
    value_font = _doc_font(24, True)
    small_font = _doc_font(21, False)

    draw.text((45, 28), "Current Annual Cost Context", font=title_font, fill=navy)
    draw.line((45, 94, width - 45, 94), fill=border, width=3)

    # Left panel: million-scale values.
    left_x0, left_x1 = 45, 710
    draw.text((left_x0, 120), "Turnover & Total Exposure", font=section_font, fill=navy)
    left_rows = [
        ("Annual turnover cost", annual_turnover_cost, navy),
        ("Total current exposure", current_cost_exposure, blue),
    ]
    left_max = max([v for _, v, _ in left_rows] + [1])
    y = 205
    for label, value, color in left_rows:
        draw.text((left_x0, y), label, font=label_font, fill=muted)
        bar_y = y + 43
        bar_width = int(520 * value / left_max)
        _draw_rounded_bar(draw, (left_x0, bar_y, left_x0 + max(bar_width, 8), bar_y + 52), color, 12)
        draw.text((left_x0 + 535, bar_y + 9), _compact_currency(value), font=value_font, fill=navy)
        y += 175

    draw.rounded_rectangle((left_x0, 595, left_x1, 742), radius=16, fill="#F4F8FF", outline=border, width=2)
    draw.text((left_x0 + 24, 620), "Average cost per employee departure", font=small_font, fill=muted)
    draw.text((left_x0 + 24, 666), _compact_currency(cost_per_departure), font=_doc_font(34, True), fill=navy)

    # Divider.
    draw.line((750, 120, 750, 745), fill=border, width=3)

    # Right panel: listening-program scale.
    right_x0 = 795
    draw.text((right_x0, 120), "Listening Program Breakdown", font=section_font, fill=navy)
    rows = [
        ("Software", software_cost, blue),
        ("Internal HR effort", internal_cost, medium),
        ("External support", external_cost, sky),
    ]
    right_max = max([v for _, v, _ in rows] + [1])
    y = 200
    for label, value, color in rows:
        draw.text((right_x0, y), label, font=label_font, fill=muted)
        bar_y = y + 40
        bar_width = int(500 * value / right_max)
        _draw_rounded_bar(draw, (right_x0, bar_y, right_x0 + max(bar_width, 8), bar_y + 48), color, 11)
        draw.text((right_x0 + 515, bar_y + 8), _compact_currency(value), font=value_font, fill=navy)
        y += 145

    draw.rounded_rectangle((right_x0, 632, width - 45, 742), radius=16, fill="#EBF4FF", outline=border, width=2)
    draw.text((right_x0 + 24, 651), "Current listening program total", font=small_font, fill=muted)
    draw.text((right_x0 + 24, 692), _compact_currency(current_listening_cost), font=_doc_font(34, True), fill=navy)

    output = BytesIO()
    image.save(output, format="PNG", optimize=True)
    output.seek(0)
    return output


def _draw_range_row(draw, *, x0, x1, y, low, high, scale_max, label, color, label_font, value_font, muted, navy):
    draw.text((x0, y), label, font=label_font, fill=muted)
    line_y = y + 43
    scale_max = max(scale_max, 1)
    start = x0 + 10 + int((x1 - x0 - 225) * max(low, 0) / scale_max)
    finish = x0 + 10 + int((x1 - x0 - 225) * max(high, 0) / scale_max)
    finish = max(finish, start + 12)
    draw.line((x0 + 10, line_y, x1 - 215, line_y), fill="#E5ECF4", width=12)
    draw.line((start, line_y, finish, line_y), fill=color, width=18)
    draw.ellipse((start - 10, line_y - 10, start + 10, line_y + 10), fill=color)
    draw.ellipse((finish - 10, line_y - 10, finish + 10, line_y + 10), fill=color)
    draw.text((x1 - 200, y + 23), _compact_range(low, high), font=value_font, fill=navy)


def _make_pythia_savings_chart(
    *,
    pythia_annual_low: float,
    pythia_annual_high: float,
    pythia_setup_low: float,
    pythia_setup_high: float,
    first_year_hsd_low: float,
    first_year_hsd_high: float,
    first_year_savings_low: float,
    first_year_savings_high: float,
    ongoing_savings_low: float,
    ongoing_savings_high: float,
) -> BytesIO:
    """Create separate-scale range charts for Pythia cost and potential savings."""
    width, height = 1500, 820
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    navy = "#1B2A4A"
    blue = "#2D6DB5"
    medium = "#4A90D9"
    green = "#07946F"
    light_green = "#35B98E"
    border = "#D8E2EE"
    muted = "#667085"

    title_font = _doc_font(44, True)
    section_font = _doc_font(29, True)
    label_font = _doc_font(23, False)
    value_font = _doc_font(23, True)
    note_font = _doc_font(19, False)

    draw.text((45, 28), "Directional Pythia Cost & Potential Savings", font=title_font, fill=navy)
    draw.line((45, 94, width - 45, 94), fill=border, width=3)

    # Costs panel.
    draw.text((45, 118), "Pythia Cost Ranges", font=section_font, fill=navy)
    cost_max = max(pythia_annual_high, pythia_setup_high, first_year_hsd_high, 1) * 1.08
    cost_rows = [
        ("Annual platform", pythia_annual_low, pythia_annual_high, blue),
        ("One-time setup", pythia_setup_low, pythia_setup_high, medium),
        ("First-year HSD cost", first_year_hsd_low, first_year_hsd_high, navy),
    ]
    y = 175
    for label, low, high, color in cost_rows:
        _draw_range_row(
            draw,
            x0=45,
            x1=1455,
            y=y,
            low=low,
            high=high,
            scale_max=cost_max,
            label=label,
            color=color,
            label_font=label_font,
            value_font=value_font,
            muted=muted,
            navy=navy,
        )
        y += 112

    draw.line((45, 516, width - 45, 516), fill=border, width=3)

    # Savings panel uses its own scale.
    draw.text((45, 540), "Potential Savings Ranges*", font=section_font, fill=navy)
    savings_max = max(first_year_savings_high, ongoing_savings_high, 1) * 1.08
    savings_rows = [
        ("First-year savings", first_year_savings_low, first_year_savings_high, green),
        ("Ongoing savings", ongoing_savings_low, ongoing_savings_high, light_green),
    ]
    y = 595
    for label, low, high, color in savings_rows:
        _draw_range_row(
            draw,
            x0=45,
            x1=1455,
            y=y,
            low=low,
            high=high,
            scale_max=savings_max,
            label=label,
            color=color,
            label_font=label_font,
            value_font=value_font,
            muted=muted,
            navy=navy,
        )
        y += 105

    draw.text((45, 775), "*Assumes all entered current listening costs are replaced or avoided.", font=note_font, fill=muted)

    output = BytesIO()
    image.save(output, format="PNG", optimize=True)
    output.seek(0)
    return output



def _make_financial_snapshot_chart(
    *,
    cost_per_departure: float,
    annual_turnover_cost: float,
    software_cost: float,
    internal_cost: float,
    external_cost: float,
    current_listening_cost: float,
    current_cost_exposure: float,
    pythia_annual_low: float,
    pythia_annual_high: float,
    pythia_setup_low: float,
    pythia_setup_high: float,
    first_year_hsd_low: float,
    first_year_hsd_high: float,
    first_year_savings_low: float,
    first_year_savings_high: float,
    ongoing_savings_low: float,
    ongoing_savings_high: float,
) -> BytesIO:
    """Create a large, full-width financial snapshot for the one-page Word report."""
    width, height = 2200, 760
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    navy = "#1B2A4A"
    blue = "#2D6DB5"
    medium = "#4A90D9"
    sky = "#BFDFFF"
    green = "#07946F"
    light_green = "#35B98E"
    muted = "#667085"
    border = "#D8E2EE"
    panel_fill = "#F8FAFC"
    track = "#E5ECF4"

    panel_title_font = _doc_font(44, True)
    label_font = _doc_font(35, False)
    value_font = _doc_font(40, True)
    note_font = _doc_font(30, False)

    def readable_range(low: float, high: float) -> str:
        if abs(low - high) < 0.01:
            return _compact_currency(low)
        low_text = _compact_currency(low)
        high_text = _compact_currency(high).replace("$", "")
        return f"{low_text} to {high_text}"

    def right_text(x_right: int, y: int, value: str, font, fill: str) -> None:
        box = draw.textbbox((0, 0), value, font=font)
        text_width = box[2] - box[0]
        draw.text((x_right - text_width, y), value, font=font, fill=fill)

    margin = 28
    gap_x = 34
    gap_y = 26
    panel_w = (width - (2 * margin) - gap_x) // 2
    panel_h = (height - (2 * margin) - gap_y) // 2

    panels = [
        (margin, margin, margin + panel_w, margin + panel_h),
        (margin + panel_w + gap_x, margin, width - margin, margin + panel_h),
        (margin, margin + panel_h + gap_y, margin + panel_w, height - margin),
        (margin + panel_w + gap_x, margin + panel_h + gap_y, width - margin, height - margin),
    ]

    for box in panels:
        draw.rounded_rectangle(box, radius=22, fill=panel_fill, outline=border, width=3)

    # Panel 1: current annual cost.
    x0, y0, x1, y1 = panels[0]
    draw.text((x0 + 30, y0 + 20), "Current Annual Cost Context", font=panel_title_font, fill=navy)
    rows = [
        ("Annual turnover cost", annual_turnover_cost, navy),
        ("Total current exposure", current_cost_exposure, blue),
    ]
    max_value = max(annual_turnover_cost, current_cost_exposure, 1)
    y = y0 + 92
    for label, value, color in rows:
        draw.text((x0 + 30, y), label, font=label_font, fill=muted)
        bar_x0 = x0 + 475
        bar_x1 = x1 - 215
        bar_width = max(bar_x1 - bar_x0, 10)
        value_width = int(bar_width * value / max_value)
        draw.rounded_rectangle((bar_x0, y + 6, bar_x0 + max(value_width, 12), y + 43), radius=12, fill=color)
        right_text(x1 - 30, y + 1, _compact_currency(value), value_font, navy)
        y += 78
    draw.text(
        (x0 + 30, y1 - 47),
        f"Average cost per departure: {_compact_currency(cost_per_departure)}",
        font=note_font,
        fill=muted,
    )

    # Panel 2: Pythia costs.
    x0, y0, x1, y1 = panels[1]
    draw.text((x0 + 30, y0 + 20), "Directional Pythia Cost", font=panel_title_font, fill=navy)
    pythia_rows = [
        ("Annual platform", pythia_annual_low, pythia_annual_high, blue),
        ("One-time setup", pythia_setup_low, pythia_setup_high, medium),
        ("First-year cost", first_year_hsd_low, first_year_hsd_high, navy),
    ]
    cost_max = max(pythia_annual_high, pythia_setup_high, first_year_hsd_high, 1)
    y = y0 + 85
    for label, low, high, color in pythia_rows:
        draw.text((x0 + 30, y), label, font=label_font, fill=muted)
        track_x0 = x0 + 350
        track_x1 = x1 - 430
        track_width = max(track_x1 - track_x0, 20)
        draw.rounded_rectangle((track_x0, y + 8, track_x1, y + 39), radius=12, fill=track)
        range_start = track_x0 + int(track_width * max(low, 0) / cost_max)
        range_end = track_x0 + int(track_width * max(high, 0) / cost_max)
        range_end = max(range_end, range_start + 20)
        draw.rounded_rectangle((range_start, y + 3, range_end, y + 44), radius=13, fill=color)
        right_text(x1 - 30, y + 1, readable_range(low, high), value_font, navy)
        y += 67

    # Panel 3: listening cost breakdown.
    x0, y0, x1, y1 = panels[2]
    draw.text((x0 + 30, y0 + 20), "Listening Program Breakdown", font=panel_title_font, fill=navy)
    categories = [
        ("Software", software_cost, blue),
        ("Internal HR effort", internal_cost, medium),
        ("External support", external_cost, sky),
    ]
    max_listening = max(software_cost, internal_cost, external_cost, 1)
    y = y0 + 85
    for label, value, color in categories:
        draw.text((x0 + 30, y), label, font=label_font, fill=muted)
        bar_x0 = x0 + 365
        bar_x1 = x1 - 205
        bar_width = max(bar_x1 - bar_x0, 10)
        value_width = int(bar_width * value / max_listening)
        draw.rounded_rectangle((bar_x0, y + 5, bar_x0 + max(value_width, 12), y + 42), radius=12, fill=color)
        right_text(x1 - 30, y + 1, _compact_currency(value), value_font, navy)
        y += 65
    draw.text(
        (x0 + 30, y1 - 47),
        f"Current listening total: {_compact_currency(current_listening_cost)}",
        font=note_font,
        fill=muted,
    )

    # Panel 4: potential savings.
    x0, y0, x1, y1 = panels[3]
    draw.text((x0 + 30, y0 + 20), "Potential Savings*", font=panel_title_font, fill=navy)
    savings_rows = [
        ("First-year savings", first_year_savings_low, first_year_savings_high, green),
        ("Ongoing savings", ongoing_savings_low, ongoing_savings_high, light_green),
    ]
    savings_max = max(first_year_savings_high, ongoing_savings_high, 1)
    y = y0 + 100
    for label, low, high, color in savings_rows:
        draw.text((x0 + 30, y), label, font=label_font, fill=muted)
        track_x0 = x0 + 430
        track_x1 = x1 - 430
        track_width = max(track_x1 - track_x0, 20)
        draw.rounded_rectangle((track_x0, y + 8, track_x1, y + 39), radius=12, fill=track)
        range_start = track_x0 + int(track_width * max(low, 0) / savings_max)
        range_end = track_x0 + int(track_width * max(high, 0) / savings_max)
        range_end = max(range_end, range_start + 20)
        draw.rounded_rectangle((range_start, y + 3, range_end, y + 44), radius=13, fill=color)
        right_text(x1 - 30, y + 1, readable_range(low, high), value_font, navy)
        y += 88
    draw.text(
        (x0 + 30, y1 - 47),
        "*Full-replacement scenario; not guaranteed.",
        font=note_font,
        fill=muted,
    )

    output = BytesIO()
    image.save(output, format="PNG", optimize=True)
    output.seek(0)
    return output

def _remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def _set_cell_margins(cell, top=40, start=50, bottom=40, end=50) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def create_hq_brief_docx(
    *,
    company: str,
    industry: str,
    employee_count_range: str,
    pythia_annual_low: float,
    pythia_annual_high: float,
    pythia_setup_low: float,
    pythia_setup_high: float,
    turnover_rate: float,
    annual_departures: int,
    cost_per_departure: float,
    software_cost: float,
    internal_cost: float,
    external_cost: float,
    current_listening_cost: float,
    annual_turnover_cost: float,
    current_cost_exposure: float,
    first_year_hsd_low: float,
    first_year_hsd_mid: float,
    first_year_hsd_high: float,
    first_year_savings_low: float,
    first_year_savings_mid: float,
    first_year_savings_high: float,
    ongoing_savings_low: float,
    ongoing_savings_mid: float,
    ongoing_savings_high: float,
    maturity_score: float,
    retention_plan: str,
) -> bytes:
    """Create a compact one-page company-specific Word brief."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.20)
    section.bottom_margin = Inches(0.20)
    section.left_margin = Inches(0.25)
    section.right_margin = Inches(0.25)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(7.6)
    styles["Normal"].paragraph_format.space_after = Pt(1)

    # Compact header: logo at top left, prospect title immediately beside it.
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(header_table)
    header_table.columns[0].width = Inches(1.15)
    header_table.columns[1].width = Inches(6.45)
    left_cell, right_cell = header_table.rows[0].cells
    _set_cell_margins(left_cell, 0, 0, 0, 70)
    _set_cell_margins(right_cell, 0, 40, 0, 0)
    left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    if HSD_LOGO_PATH.exists():
        logo_paragraph = left_cell.paragraphs[0]
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(str(HSD_LOGO_PATH), width=Inches(1.05))

    title_paragraph = right_cell.paragraphs[0]
    title_paragraph.paragraph_format.space_after = Pt(1)
    title_run = title_paragraph.add_run(f"{company} Employee Listening Enhancement Brief")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(17)
    title_run.font.color.rgb = RGBColor(27, 42, 74)

    subtitle_paragraph = right_cell.add_paragraph()
    subtitle_paragraph.paragraph_format.space_after = Pt(0)
    subtitle_run = subtitle_paragraph.add_run(
        f"Prepared using the HSD HQ Brief directional pre-sales model | {date.today().strftime('%B %d, %Y')}"
    )
    subtitle_run.italic = True
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(7.5)
    subtitle_run.font.color.rgb = RGBColor(107, 114, 128)

    # Executive overview.
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(2)
    heading.paragraph_format.space_after = Pt(1)
    run = heading.add_run("Executive overview")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11.5)
    run.font.color.rgb = RGBColor(27, 42, 74)

    overview = doc.add_paragraph()
    overview.paragraph_format.space_after = Pt(2)
    overview_run = overview.add_run(
        f"This brief summarizes the cost information entered for {company} and compares the current "
        "employee-listening program cost with the directional Pythia platform and setup estimates. "
        "Potential savings use a full-replacement scenario and are not guaranteed."
    )
    overview_run.font.name = "Arial"
    overview_run.font.size = Pt(7.7)

    # Prospect profile remains as a compact table.
    profile_heading = doc.add_paragraph()
    profile_heading.paragraph_format.space_before = Pt(1)
    profile_heading.paragraph_format.space_after = Pt(1)
    profile_run = profile_heading.add_run("Prospect profile")
    profile_run.bold = True
    profile_run.font.name = "Arial"
    profile_run.font.size = Pt(11.5)
    profile_run.font.color.rgb = RGBColor(27, 42, 74)

    profile_table = doc.add_table(rows=3, cols=4)
    profile_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    profile_table.style = "Table Grid"
    profile_rows = [
        ("Industry", industry or "Not entered", "Employee range", employee_count_range),
        ("Turnover rate", percent(turnover_rate), "Annual departures", f"{annual_departures:,}"),
        ("Listening maturity", f"{maturity_score:.0f}/100", "Retention plan", retention_plan),
    ]
    for row, values in zip(profile_table.rows, profile_rows):
        for index, value in enumerate(values):
            is_label = index in (0, 2)
            set_cell_text(row.cells[index], value, bold=is_label, color="1B2A4A" if is_label else None)
            row.cells[index].paragraphs[0].runs[0].font.size = Pt(7.5)
            _set_cell_margins(row.cells[index], 20, 45, 20, 45)
            if is_label:
                shade_cell(row.cells[index], "EBF4FF")

    # Replace the two large financial tables with two graphs.
    financial_heading = doc.add_paragraph()
    financial_heading.paragraph_format.space_before = Pt(2)
    financial_heading.paragraph_format.space_after = Pt(1)
    financial_run = financial_heading.add_run("Financial snapshot")
    financial_run.bold = True
    financial_run.font.name = "Arial"
    financial_run.font.size = Pt(11.5)
    financial_run.font.color.rgb = RGBColor(27, 42, 74)

    financial_chart = _make_financial_snapshot_chart(
        cost_per_departure=cost_per_departure,
        annual_turnover_cost=annual_turnover_cost,
        software_cost=software_cost,
        internal_cost=internal_cost,
        external_cost=external_cost,
        current_listening_cost=current_listening_cost,
        current_cost_exposure=current_cost_exposure,
        pythia_annual_low=pythia_annual_low,
        pythia_annual_high=pythia_annual_high,
        pythia_setup_low=pythia_setup_low,
        pythia_setup_high=pythia_setup_high,
        first_year_hsd_low=first_year_hsd_low,
        first_year_hsd_high=first_year_hsd_high,
        first_year_savings_low=first_year_savings_low,
        first_year_savings_high=first_year_savings_high,
        ongoing_savings_low=ongoing_savings_low,
        ongoing_savings_high=ongoing_savings_high,
    )
    financial_paragraph = doc.add_paragraph()
    financial_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    financial_paragraph.paragraph_format.space_after = Pt(1)
    financial_paragraph.add_run().add_picture(financial_chart, width=Inches(7.95))

    # Keep the sales message, but make it concise.
    sales = doc.add_paragraph()
    sales.paragraph_format.space_before = Pt(1)
    sales.paragraph_format.space_after = Pt(2)
    sales_run = sales.add_run(
        f"Sales message: Current employee-listening costs are {money(current_listening_cost)}. "
        f"Estimated first-year Pythia cost is "
        f"{money_range(first_year_hsd_low, first_year_hsd_high) if first_year_hsd_high > 0 else 'not selected'}, "
        f"with directional first-year savings of "
        f"{signed_money_range(first_year_savings_low, first_year_savings_high) if first_year_hsd_high > 0 and current_listening_cost > 0 else 'not available'}."
    )
    sales_run.bold = True
    sales_run.font.name = "Arial"
    sales_run.font.size = Pt(7.6)
    sales_run.font.color.rgb = RGBColor(27, 42, 74)

    # Keep savings scenario detail as a compact table on the same page.
    scenario_heading = doc.add_paragraph()
    scenario_heading.paragraph_format.space_before = Pt(1)
    scenario_heading.paragraph_format.space_after = Pt(1)
    scenario_run = scenario_heading.add_run("Savings scenario detail")
    scenario_run.bold = True
    scenario_run.font.name = "Arial"
    scenario_run.font.size = Pt(11.5)
    scenario_run.font.color.rgb = RGBColor(27, 42, 74)

    scenario_table = doc.add_table(rows=4, cols=5)
    scenario_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    scenario_table.style = "Table Grid"
    headers = ["Scenario", "First-year HSD cost", "First-year savings", "Ongoing annual cost", "Ongoing savings"]
    for index, header in enumerate(headers):
        set_cell_text(scenario_table.rows[0].cells[index], header, bold=True, color="FFFFFF")
        scenario_table.rows[0].cells[index].paragraphs[0].runs[0].font.size = Pt(7)
        shade_cell(scenario_table.rows[0].cells[index], "1B2A4A")
        _set_cell_margins(scenario_table.rows[0].cells[index], 18, 30, 18, 30)

    scenarios = [
        ("Low cost", first_year_hsd_low, first_year_savings_high, pythia_annual_low, ongoing_savings_high),
        ("Midpoint", first_year_hsd_mid, first_year_savings_mid, (pythia_annual_low + pythia_annual_high) / 2, ongoing_savings_mid),
        ("High cost", first_year_hsd_high, first_year_savings_low, pythia_annual_high, ongoing_savings_low),
    ]
    for row, values in zip(scenario_table.rows[1:], scenarios):
        display_values = [values[0], money(values[1]), signed_money_value(values[2]), money(values[3]), signed_money_value(values[4])]
        for index, value in enumerate(display_values):
            set_cell_text(row.cells[index], value, bold=index == 0)
            row.cells[index].paragraphs[0].runs[0].font.size = Pt(7)
            _set_cell_margins(row.cells[index], 16, 30, 16, 30)
            if index == 0:
                shade_cell(row.cells[index], "EBF4FF")

    # Recommended approach remains; calculation methodology is intentionally removed.
    recommendation_heading = doc.add_paragraph()
    recommendation_heading.paragraph_format.space_before = Pt(2)
    recommendation_heading.paragraph_format.space_after = Pt(0)
    recommendation_run = recommendation_heading.add_run("Recommended HSD approach")
    recommendation_run.bold = True
    recommendation_run.font.name = "Arial"
    recommendation_run.font.size = Pt(11)
    recommendation_run.font.color.rgb = RGBColor(27, 42, 74)

    recommendations = [
        "Validate which software, HR effort, and external support costs would actually be reduced or removed.",
        "Confirm final HSD scope and proposal price before presenting savings externally.",
        "Treat first-year savings separately because the setup fee occurs only once.",
        "Do not include turnover reduction until HSD approves an impact assumption.",
    ]
    for item in recommendations:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.16)
        paragraph.paragraph_format.first_line_indent = Inches(-0.12)
        paragraph.paragraph_format.space_after = Pt(0)
        bullet_run = paragraph.add_run(item)
        bullet_run.font.name = "Arial"
        bullet_run.font.size = Pt(7.2)

    interpretation_heading = doc.add_paragraph()
    interpretation_heading.paragraph_format.space_before = Pt(1)
    interpretation_heading.paragraph_format.space_after = Pt(0)
    interpretation_run = interpretation_heading.add_run("Important interpretation")
    interpretation_run.bold = True
    interpretation_run.font.name = "Arial"
    interpretation_run.font.size = Pt(10.5)
    interpretation_run.font.color.rgb = RGBColor(27, 42, 74)

    disclaimer = doc.add_paragraph()
    disclaimer.paragraph_format.space_after = Pt(0)
    disclaimer_run = disclaimer.add_run(
        "Potential savings are directional and depend on whether the client can eliminate or reduce the entered software, "
        "internal HR effort, and external support costs. Turnover reduction, ROI, and productivity gains are not included. "
        "Final pricing, scope, and client-facing claims should be validated by HSD leadership."
    )
    disclaimer_run.font.name = "Arial"
    disclaimer_run.font.size = Pt(7.1)
    disclaimer_run.font.color.rgb = RGBColor(75, 85, 99)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(1)
    footer.paragraph_format.space_after = Pt(0)
    footer_run = footer.add_run(f"HSD Metrics | {HSD_WEBSITE}")
    footer_run.bold = True
    footer_run.font.name = "Arial"
    footer_run.font.size = Pt(7.2)
    footer_run.font.color.rgb = RGBColor(45, 109, 181)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()

# --------------------------------------------------
# SIDEBAR INPUTS
# --------------------------------------------------
st.sidebar.title("Build Prospect Profile")
company = st.sidebar.text_input("Company Name", "New Prospect Company")
industry = st.sidebar.text_input("Industry", placeholder="e.g., Insurance")

employee_count_range = st.sidebar.selectbox(
    "Employee Count Range",
    ["Select employee range", *PYTHIA_ESTIMATED_PRICING.keys()],
    index=0,
    help="Select the prospect's approximate total employee-count range.",
)

pythia_selected = employee_count_range != "Select employee range"
if pythia_selected:
    pythia_annual_low, pythia_annual_high = PYTHIA_ESTIMATED_PRICING[employee_count_range]["annual"]
    pythia_setup_low, pythia_setup_high = PYTHIA_ESTIMATED_PRICING[employee_count_range]["setup"]
else:
    pythia_annual_low = pythia_annual_high = 0.0
    pythia_setup_low = pythia_setup_high = 0.0

st.sidebar.markdown("---")
st.sidebar.subheader("Turnover & Cost Inputs")

turnover_rate = st.sidebar.number_input(
    "Estimated Turnover Rate %",
    min_value=0.0,
    max_value=100.0,
    value=0.0,
    step=0.5,
    help="This is shown for context. Annual departures are entered separately because headcount is not used.",
)

annual_departures = st.sidebar.number_input(
    "Annual Employee Departures",
    min_value=0,
    value=0,
    step=1,
    help="Enter the number of employees who leave during a typical year.",
)

cost_per_departure = st.sidebar.number_input(
    "Average Cost per Employee Departure",
    min_value=0.0,
    value=0.0,
    step=1000.0,
    help="This may include recruiting, onboarding, vacancy time, training, lost productivity, and knowledge loss.",
)

software_cost = st.sidebar.number_input(
    "Current Annual Listening Software Cost",
    min_value=0.0,
    value=0.0,
    step=5000.0,
)
internal_cost = st.sidebar.number_input(
    "Current Internal HR Effort Cost",
    min_value=0.0,
    value=0.0,
    step=5000.0,
)
external_cost = st.sidebar.number_input(
    "Current External Support Cost",
    min_value=0.0,
    value=0.0,
    step=5000.0,
)

st.sidebar.markdown("---")
st.sidebar.subheader("Listening Inputs")
maturity_score = st.sidebar.number_input(
    "Current Listening Maturity Score",
    min_value=0.0,
    max_value=100.0,
    value=0.0,
    step=1.0,
    help="Use this only when the score comes from an agreed assessment or questionnaire.",
)
retention_plan = st.sidebar.radio(
    "Retention Action Plan",
    ["Not entered", "No", "In progress", "Yes"],
    index=0,
    help="Select the option that best describes the prospect's current retention planning status.",
)

# --------------------------------------------------
# CALCULATIONS
# --------------------------------------------------
annual_turnover_cost = annual_departures * cost_per_departure
current_listening_cost = software_cost + internal_cost + external_cost
current_cost_exposure = annual_turnover_cost + current_listening_cost

pythia_annual_mid = (pythia_annual_low + pythia_annual_high) / 2
pythia_setup_mid = (pythia_setup_low + pythia_setup_high) / 2

first_year_hsd_low = pythia_annual_low + pythia_setup_low
first_year_hsd_mid = pythia_annual_mid + pythia_setup_mid
first_year_hsd_high = pythia_annual_high + pythia_setup_high

# Conservative savings use the higher HSD cost; optimistic savings use the lower HSD cost.
first_year_savings_low = current_listening_cost - first_year_hsd_high
first_year_savings_mid = current_listening_cost - first_year_hsd_mid
first_year_savings_high = current_listening_cost - first_year_hsd_low

ongoing_savings_low = current_listening_cost - pythia_annual_high
ongoing_savings_mid = current_listening_cost - pythia_annual_mid
ongoing_savings_high = current_listening_cost - pythia_annual_low

savings_available = pythia_selected and current_listening_cost > 0

savings_scenarios = pd.DataFrame(
    {
        "Scenario": ["Low", "Midpoint", "High"],
        "Potential First-Year Savings": [
            first_year_savings_low,
            first_year_savings_mid,
            first_year_savings_high,
        ],
    }
)
savings_scenarios["Display Label"] = savings_scenarios["Potential First-Year Savings"].map(signed_money_value)

company_html = html.escape(company or "New Prospect Company")
industry_html = html.escape(industry or "Not entered")
employee_range_html = html.escape(employee_count_range if pythia_selected else "Not selected")

# --------------------------------------------------
# HEADER
# --------------------------------------------------
st.markdown(
    f"""
    <div class="hsd-header">
        <div class="hsd-logo-wrap">
            <a href="{HSD_WEBSITE}" target="_blank" rel="noopener noreferrer">
                <img src="{HSD_LOGO_SRC}" alt="HSD Metrics logo" title="Open HSD Metrics website">
            </a>
        </div>
        <div>
            <h1>HSD HQ Brief</h1>
            <p>A directional pre-sales tool for comparing current listening spend with Pythia pricing and potential savings.</p>
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)

st.caption("Click the HSD logo to open the HSD Metrics website. Results use only the values entered in the sidebar.")

st.markdown(
    textwrap.dedent(
        """
        <div class="hsd-results">
            <div class="hsd-results-title">Proven Results from Our Clients</div>
            <div class="hsd-results-grid">
                <div class="hsd-result-stat">
                    <div class="hsd-result-icon hsd-icon-blue">👥</div>
                    <div>
                        <div class="hsd-result-value hsd-value-blue">97%</div>
                        <div class="hsd-result-label">Client retention rate</div>
                    </div>
                </div>
                <div class="hsd-result-stat">
                    <div class="hsd-result-icon hsd-icon-green">↘</div>
                    <div>
                        <div class="hsd-result-value hsd-value-green">80%</div>
                        <div class="hsd-result-label">Of clients report a decrease in turnover</div>
                    </div>
                </div>
                <div class="hsd-result-stat">
                    <div class="hsd-result-icon hsd-icon-purple">💬</div>
                    <div>
                        <div class="hsd-result-value hsd-value-purple">4M+</div>
                        <div class="hsd-result-label">Employees surveyed across industries</div>
                    </div>
                </div>
                <div class="hsd-result-quote">
                    <div class="hsd-quote-mark">“</div>
                    <div class="hsd-quote-text">
                        We had tried two other platforms before HSD. Our response rates went
                        from 22% to over 80% in the first cycle.
                    </div>
                    <div class="hsd-quote-source">
                        — Chief People Officer,<br>Healthcare Organization
                    </div>
                </div>
            </div>
        </div>
        """
    ),
    unsafe_allow_html=True,
)

# --------------------------------------------------
# CLEANER DASHBOARD TABS
# --------------------------------------------------
tab1, tab2, tab3 = st.tabs(
    [
        "Executive Summary",
        "Cost Details",
        "HQ Brief Summary",
    ]
)

# --------------------------------------------------
# DYNAMIC INTERPRETATION TEXT
# --------------------------------------------------
if not pythia_selected:
    savings_sentence = "Select an employee-count range to calculate the Pythia cost and potential savings."
elif current_listening_cost <= 0:
    savings_sentence = "Enter the client’s current listening-program costs to calculate potential savings."
elif first_year_savings_high < 0:
    savings_sentence = (
        f"The estimated first-year Pythia cost is {signed_money_range(abs(first_year_savings_high), abs(first_year_savings_low))} "
        "above the entered current listening-program cost."
    )
elif first_year_savings_low >= 0:
    savings_sentence = (
        f"Potential first-year savings are {signed_money_range(first_year_savings_low, first_year_savings_high)} "
        "under a full-replacement scenario."
    )
else:
    savings_sentence = (
        "The first-year result ranges from a possible additional cost to possible savings, depending on the final Pythia price."
    )

pythia_sentence = (
    f"For the selected employee range of {employee_range_html}, the directional annual Pythia estimate is "
    f"{money_range(pythia_annual_low, pythia_annual_high)}, with a one-time setup estimate of "
    f"{money_range(pythia_setup_low, pythia_setup_high)}."
    if pythia_selected
    else "Select an employee-count range to generate the directional Pythia platform and setup estimates."
)

# --------------------------------------------------
# TAB 1: EXECUTIVE SUMMARY
# --------------------------------------------------
with tab1:
    st.header("Executive Summary")

    st.markdown(
        f"""
        <div class="hsd-prospect-line">
            <strong>{company_html}</strong>
            &nbsp;&nbsp;|&nbsp;&nbsp;
            {industry_html}
            &nbsp;&nbsp;|&nbsp;&nbsp;
            Employee range: {employee_range_html}
        </div>
        """,
        unsafe_allow_html=True,
    )

    kpi1, kpi2, kpi3 = st.columns(3)
    kpi1.metric("Annual Turnover Cost", compact_money(annual_turnover_cost))
    kpi2.metric("Current Listening Program Cost", compact_money(current_listening_cost))
    kpi3.metric(
        "Potential First-Year Savings*",
        metric_signed_money_range(first_year_savings_low, first_year_savings_high)
        if savings_available
        else "Not available",
    )

    visual_left, visual_right = st.columns([1.15, 1])

    with visual_left:
        comparison_data = pd.DataFrame(
            {
                "Cost": [
                    "Current Listening Program",
                    "Pythia First-Year Low",
                    "Pythia First-Year Midpoint",
                    "Pythia First-Year High",
                ],
                "Annual Amount": [
                    current_listening_cost,
                    first_year_hsd_low,
                    first_year_hsd_mid,
                    first_year_hsd_high,
                ],
            }
        )

        if pythia_selected and comparison_data["Annual Amount"].sum() > 0:
            fig_summary_comparison = px.bar(
                comparison_data,
                x="Cost",
                y="Annual Amount",
                text="Annual Amount",
                title="Current Listening Cost vs Estimated First-Year Pythia Cost",
                color="Cost",
                color_discrete_sequence=[
                    HSD_NAVY,
                    HSD_SKY_BLUE,
                    HSD_MEDIUM_BLUE,
                    HSD_BLUE,
                ],
            )
            fig_summary_comparison.update_traces(
                texttemplate="$%{text:,.0f}",
                textposition="outside",
            )
            fig_summary_comparison.update_layout(
                xaxis_title="",
                yaxis_title="Annual Cost",
                showlegend=False,
            )
            fig_summary_comparison = apply_hsd_theme(fig_summary_comparison)
            st.plotly_chart(fig_summary_comparison, use_container_width=True)
        else:
            st.info("Select an employee range and enter current listening costs to display the comparison.")

    with visual_right:
        st.markdown(
            f"""
            <div class="hsd-pricing-box">
                <h3>Pythia Cost & Savings Estimate</h3>
                <div class="hsd-pricing-row">
                    <span class="hsd-pricing-label">Employee count range</span>
                    <span class="hsd-pricing-value">{employee_range_html}</span>
                </div>
                <div class="hsd-pricing-row">
                    <span class="hsd-pricing-label">Annual platform estimate</span>
                    <span class="hsd-pricing-value">
                        {metric_money_range(pythia_annual_low, pythia_annual_high) if pythia_selected else "Not selected"}
                    </span>
                </div>
                <div class="hsd-pricing-row">
                    <span class="hsd-pricing-label">One-time setup estimate</span>
                    <span class="hsd-pricing-value">
                        {metric_money_range(pythia_setup_low, pythia_setup_high) if pythia_selected else "Not selected"}
                    </span>
                </div>
                <div class="hsd-pricing-row">
                    <span class="hsd-pricing-label">First-year Pythia cost</span>
                    <span class="hsd-pricing-value">
                        {metric_money_range(first_year_hsd_low, first_year_hsd_high) if pythia_selected else "Not selected"}
                    </span>
                </div>
                <div class="hsd-pricing-row">
                    <span class="hsd-pricing-label">Potential ongoing annual savings*</span>
                    <span class="hsd-pricing-value">
                        {metric_signed_money_range(ongoing_savings_low, ongoing_savings_high) if savings_available else "Not available"}
                    </span>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    st.markdown(
        f"""
        <div class="hsd-insight-box">
            <h3>What the entered information suggests</h3>
            <p>
                <b>Turnover exposure:</b> {annual_departures:,} annual departures at
                {money(cost_per_departure)} per departure produce an estimated annual
                turnover cost of <b>{money(annual_turnover_cost)}</b>.
            </p>
            <p>
                <b>Current listening spend:</b> The entered software, internal HR effort,
                and external support costs total <b>{money(current_listening_cost)}</b> annually.
            </p>
            <p><b>Potential savings:</b> {savings_sentence}</p>
            <p><b>Pythia estimate:</b> {pythia_sentence}</p>
            <p>
                <b>Important assumption:</b> Potential savings assume that all entered current
                listening costs can be removed or avoided after adopting HSD. Validate this with
                the client before using the figure in a proposal.
            </p>
            <p>
                <b>Listening context:</b> Current maturity is {maturity_score:.0f}/100,
                and the retention action plan is marked as “{html.escape(retention_plan)}.”
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

# --------------------------------------------------
# TAB 2: COST DETAILS
# --------------------------------------------------
with tab2:
    st.header("Cost Details")

    detail_left, detail_right = st.columns(2)

    with detail_left:
        listening_breakdown = pd.DataFrame(
            {
                "Cost Category": ["Software", "Internal HR Effort", "External Support"],
                "Amount": [software_cost, internal_cost, external_cost],
            }
        )

        if listening_breakdown["Amount"].sum() > 0:
            fig_breakdown = px.pie(
                listening_breakdown,
                names="Cost Category",
                values="Amount",
                hole=0.55,
                title="Current Listening Program Cost Breakdown",
                color="Cost Category",
                color_discrete_sequence=[HSD_BLUE, HSD_MEDIUM_BLUE, HSD_SKY_BLUE],
            )
            fig_breakdown.update_traces(
                textposition="inside",
                texttemplate="$%{value:,.0f}<br>%{percent}",
                hovertemplate="<b>%{label}</b><br>Annual cost: $%{value:,.0f}<br>Share: %{percent}<extra></extra>",
                sort=False,
            )
            fig_breakdown.add_annotation(
                text=f"<b>{money(current_listening_cost)}</b><br>Total",
                x=0.5,
                y=0.5,
                showarrow=False,
                font=dict(size=16, color=HSD_NAVY),
            )
            fig_breakdown = apply_hsd_theme(fig_breakdown)
            st.plotly_chart(fig_breakdown, use_container_width=True)
        else:
            st.info("Enter software, HR effort, or external-support costs to display the donut chart.")

    with detail_right:
        if savings_available:
            fig_savings = px.bar(
                savings_scenarios,
                x="Scenario",
                y="Potential First-Year Savings",
                text="Display Label",
                title="Potential First-Year Savings Scenarios",
                color="Scenario",
                color_discrete_sequence=[HSD_SKY_BLUE, HSD_MEDIUM_BLUE, HSD_NAVY],
            )
            fig_savings.update_traces(textposition="outside")
            fig_savings.update_layout(
                xaxis_title="",
                yaxis_title="Savings / (Additional Cost)",
                showlegend=False,
            )
            fig_savings = apply_hsd_theme(fig_savings)
            st.plotly_chart(fig_savings, use_container_width=True)
        else:
            st.info("Select an employee range and enter current listening costs to display savings scenarios.")

    exact_values = pd.DataFrame(
        {
            "Measure": [
                "Estimated turnover rate",
                "Annual employee departures",
                "Average cost per employee departure",
                "Annual turnover cost",
                "Current listening software cost",
                "Current internal HR effort cost",
                "Current external support cost",
                "Current listening program cost",
                "Total current cost exposure",
                "Pythia annual platform estimate",
                "Pythia one-time setup estimate",
                "Estimated first-year Pythia cost",
                "Potential first-year savings",
                "Potential ongoing annual savings",
            ],
            "Value": [
                percent(turnover_rate),
                f"{annual_departures:,}",
                money(cost_per_departure),
                money(annual_turnover_cost),
                money(software_cost),
                money(internal_cost),
                money(external_cost),
                money(current_listening_cost),
                money(current_cost_exposure),
                money_range(pythia_annual_low, pythia_annual_high) if pythia_selected else "Not selected",
                money_range(pythia_setup_low, pythia_setup_high) if pythia_selected else "Not selected",
                money_range(first_year_hsd_low, first_year_hsd_high) if pythia_selected else "Not selected",
                signed_money_range(first_year_savings_low, first_year_savings_high) if savings_available else "Not available",
                signed_money_range(ongoing_savings_low, ongoing_savings_high) if savings_available else "Not available",
            ],
        }
    )

    with st.expander("View exact values"):
        st.dataframe(exact_values, use_container_width=True, hide_index=True)

    with st.expander("View calculation formulas"):
        st.markdown(
            """
            **Current listening program cost** = Software cost + Internal HR effort cost + External support cost  
            **Estimated first-year Pythia cost** = Annual platform estimate + One-time setup estimate  
            **Potential first-year savings** = Current listening program cost − Estimated first-year Pythia cost  
            **Potential ongoing annual savings** = Current listening program cost − Annual Pythia platform estimate
            """
        )

    st.warning(
        "Potential savings assume that all entered current listening costs are replaced or avoided by HSD. "
        "The calculation does not include turnover reduction or productivity benefits. Validate the assumption with the client and HSD leadership."
    )

# --------------------------------------------------
# TAB 3: HQ BRIEF SUMMARY + DOWNLOAD
# --------------------------------------------------
with tab3:
    st.header("HQ Brief Summary")

    st.markdown(
        f"""
        <div class="hsd-blue-box">
            <h3>Recommended Sales Summary</h3>
            <p>
                <b>{company_html}</b> is a prospect in the <b>{industry_html}</b> industry,
                with a selected employee range of <b>{employee_range_html}</b>.
            </p>
            <p>
                Based on the entered values, estimated annual turnover cost is
                <b>{money(annual_turnover_cost)}</b>, while the current employee-listening
                program cost is <b>{money(current_listening_cost)}</b>.
            </p>
            <p>{pythia_sentence}</p>
            <p>
                <b>Potential first-year savings:</b>
                {signed_money_range(first_year_savings_low, first_year_savings_high) if savings_available else "Not available"}.
                {savings_sentence}
            </p>
            <p>
                The savings estimate assumes that all entered current listening costs can be
                eliminated or avoided. Confirm the final scope and replacement assumptions before
                including the estimate in a formal proposal.
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    docx_bytes = create_hq_brief_docx(
        company=company or "New Prospect Company",
        industry=industry,
        employee_count_range=employee_count_range if pythia_selected else "Not selected",
        pythia_annual_low=pythia_annual_low,
        pythia_annual_high=pythia_annual_high,
        pythia_setup_low=pythia_setup_low,
        pythia_setup_high=pythia_setup_high,
        turnover_rate=turnover_rate,
        annual_departures=int(annual_departures),
        cost_per_departure=cost_per_departure,
        software_cost=software_cost,
        internal_cost=internal_cost,
        external_cost=external_cost,
        current_listening_cost=current_listening_cost,
        annual_turnover_cost=annual_turnover_cost,
        current_cost_exposure=current_cost_exposure,
        first_year_hsd_low=first_year_hsd_low,
        first_year_hsd_mid=first_year_hsd_mid,
        first_year_hsd_high=first_year_hsd_high,
        first_year_savings_low=first_year_savings_low,
        first_year_savings_mid=first_year_savings_mid,
        first_year_savings_high=first_year_savings_high,
        ongoing_savings_low=ongoing_savings_low,
        ongoing_savings_mid=ongoing_savings_mid,
        ongoing_savings_high=ongoing_savings_high,
        maturity_score=maturity_score,
        retention_plan=retention_plan,
    )

    st.download_button(
        label="Download Report",
        data=docx_bytes,
        file_name=f"{safe_filename(company)}_Employee_Listening_Enhancement_Brief.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

    st.caption(
        "Directional pre-sales estimate. Final pricing, scope, replacement assumptions, and client-facing statements should be validated by HSD leadership."
    )
